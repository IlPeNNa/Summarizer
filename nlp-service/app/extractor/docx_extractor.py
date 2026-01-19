"""
Estrattore di testo da file DOCX (Microsoft Word).
"""
from typing import Optional
from docx import Document
from io import BytesIO


def extract_from_docx(file_content: bytes) -> str:
    """
    Estrae il testo da un file DOCX.
    
    Args:
        file_content: Contenuto binario del file DOCX
        
    Returns:
        Testo estratto dal documento
    """
    try:
        docx_file = BytesIO(file_content)
        doc = Document(docx_file)
        
        # Estrae il testo da tutti i paragrafi
        paragraphs = [paragraph.text for paragraph in doc.paragraphs if paragraph.text.strip()]
        
        # Estrae il testo dalle tabelle (opzionale)
        tables_text = []
        for table in doc.tables:
            for row in table.rows:
                row_text = [cell.text for cell in row.cells if cell.text.strip()]
                if row_text:
                    tables_text.append(" | ".join(row_text))
        
        # Combina paragrafi e tabelle
        all_text = paragraphs + tables_text
        full_text = "\n\n".join(all_text)
        
        return full_text
    
    except Exception as e:
        raise Exception(f"Errore nell'estrazione del DOCX: {str(e)}")


def extract_from_docx_file(file_path: str) -> str:
    """
    Estrae il testo da un file DOCX dato il percorso.
    
    Args:
        file_path: Percorso del file DOCX
        
    Returns:
        Testo estratto dal documento
    """
    try:
        with open(file_path, 'rb') as file:
            return extract_from_docx(file.read())
    except FileNotFoundError:
        raise Exception(f"File DOCX non trovato: {file_path}")
    except Exception as e:
        raise Exception(f"Errore nella lettura del file DOCX: {str(e)}")
